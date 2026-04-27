from __future__ import annotations

import asyncio
import json
import mimetypes
import os
import re
import uuid
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from html.parser import HTMLParser
from pathlib import Path
from posixpath import dirname, normpath
from typing import Any
from xml.etree import ElementTree
from zipfile import BadZipFile, ZipFile


TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".json",
    ".csv",
    ".tsv",
    ".html",
    ".htm",
    ".xml",
    ".yml",
    ".yaml",
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".sql",
    ".log",
}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg"}
MARKDOWN_EXTENSIONS = {".md", ".markdown"}


@dataclass
class ExtractedSegment:
    content: str
    metadata: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return {"content": self.content, "metadata": dict(self.metadata)}

    @classmethod
    def from_dict(cls, payload: dict[str, Any] | None) -> "ExtractedSegment":
        data = payload if isinstance(payload, dict) else {}
        return cls(
            content=str(data.get("content") or ""),
            metadata=data.get("metadata")
            if isinstance(data.get("metadata"), dict)
            else {},
        )


@dataclass
class ExtractedDocument:
    text: str
    segments: list[ExtractedSegment] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return {
            "text": self.text,
            "segments": [segment.to_dict() for segment in self.segments],
            "metadata": dict(self.metadata),
        }

    @classmethod
    def from_dict(cls, payload: dict[str, Any] | None) -> "ExtractedDocument":
        data = payload if isinstance(payload, dict) else {}
        raw_segments = data.get("segments")
        segments = raw_segments if isinstance(raw_segments, list) else []
        return cls(
            text=str(data.get("text") or ""),
            segments=[
                ExtractedSegment.from_dict(item)
                for item in segments
                if isinstance(item, dict)
            ],
            metadata=data.get("metadata")
            if isinstance(data.get("metadata"), dict)
            else {},
        )


def _utcnow_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _safe_filename(filename: str | None) -> str:
    candidate = Path((filename or "").strip() or "upload.bin").name
    return candidate or "upload.bin"


def _guess_content_type(filename: str, content_type: str | None) -> str:
    normalized = (content_type or "").strip()
    if normalized:
        return normalized
    guessed, _ = mimetypes.guess_type(filename)
    return guessed or "application/octet-stream"


def _guess_kind(filename: str, content_type: str) -> str:
    suffix = Path(filename).suffix.lower()
    normalized_type = (content_type or "").lower()
    if normalized_type.startswith("image/") or suffix in IMAGE_EXTENSIONS:
        return "image"
    if normalized_type.startswith("text/") or suffix in TEXT_EXTENSIONS:
        return "text"
    if suffix in {".pdf", ".docx", ".epub"}:
        return "document"
    return "binary"


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _base_document_metadata(path: Path, file_format: str | None = None) -> dict[str, Any]:
    suffix = path.suffix.lower().lstrip(".")
    return {
        "source_label": "file",
        "source_filename": path.name,
        "file_format": file_format or suffix or "text",
    }


def _segment(
    content: str,
    metadata: dict[str, Any],
) -> ExtractedSegment | None:
    cleaned = content.strip()
    if not cleaned:
        return None
    return ExtractedSegment(content=cleaned, metadata=dict(metadata))


def _single_segment_document(
    text: str,
    metadata: dict[str, Any],
) -> ExtractedDocument:
    segment = _segment(text, metadata)
    return ExtractedDocument(
        text=text.strip(),
        segments=[segment] if segment else [],
        metadata=dict(metadata),
    )


def _heading_metadata(
    base_metadata: dict[str, Any],
    section_path: list[str],
) -> dict[str, Any]:
    metadata = dict(base_metadata)
    if section_path:
        metadata["section_path"] = list(section_path)
        metadata["chapter_title"] = section_path[-1]
    return metadata


def _split_markdown_document(text: str, base_metadata: dict[str, Any]) -> ExtractedDocument:
    heading_pattern = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
    segments: list[ExtractedSegment] = []
    current_lines: list[str] = []
    section_path: list[str] = []
    current_section_path: list[str] = []

    def flush() -> None:
        nonlocal current_lines, current_section_path
        segment = _segment(
            "\n".join(current_lines),
            _heading_metadata(base_metadata, current_section_path),
        )
        if segment:
            segments.append(segment)
        current_lines = []

    for line in text.splitlines():
        match = heading_pattern.match(line.strip())
        if match:
            flush()
            level = len(match.group(1))
            title = match.group(2).strip().strip("#").strip()
            section_path = section_path[: level - 1] + [title]
            current_section_path = list(section_path)
            current_lines = [line]
            continue
        current_lines.append(line)

    flush()
    if not segments:
        return _single_segment_document(text, base_metadata)
    metadata = dict(base_metadata)
    metadata["section_count"] = len(segments)
    return ExtractedDocument(text=text.strip(), segments=segments, metadata=metadata)


class _HTMLTextExtractor(HTMLParser):
    BLOCK_TAGS = {
        "address",
        "article",
        "aside",
        "blockquote",
        "br",
        "div",
        "figcaption",
        "figure",
        "footer",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "header",
        "li",
        "main",
        "nav",
        "ol",
        "p",
        "pre",
        "section",
        "table",
        "td",
        "th",
        "tr",
        "ul",
    }

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._parts: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        del attrs
        normalized = tag.lower()
        if normalized in {"script", "style", "noscript"}:
            self._skip_depth += 1
        if normalized in self.BLOCK_TAGS:
            self._parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        normalized = tag.lower()
        if normalized in {"script", "style", "noscript"} and self._skip_depth:
            self._skip_depth -= 1
        if normalized in self.BLOCK_TAGS:
            self._parts.append("\n")

    def handle_data(self, data: str) -> None:
        if self._skip_depth:
            return
        normalized = " ".join(data.split())
        if normalized:
            self._parts.append(normalized)
            self._parts.append(" ")

    def text(self) -> str:
        lines = [" ".join(line.split()) for line in "".join(self._parts).splitlines()]
        return "\n".join(line for line in lines if line).strip()


def _strip_html_text(payload: bytes) -> str:
    parser = _HTMLTextExtractor()
    parser.feed(_decode_text(payload))
    parser.close()
    return parser.text()


def _xml_local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1].lower()


def _epub_rootfile_path(archive: ZipFile) -> str:
    try:
        container_xml = archive.read("META-INF/container.xml")
    except KeyError as exc:
        raise RuntimeError("EPUB is missing META-INF/container.xml") from exc
    root = ElementTree.fromstring(container_xml)
    for element in root.iter():
        if _xml_local_name(element.tag) == "rootfile":
            full_path = (element.attrib.get("full-path") or "").strip()
            if full_path:
                return full_path
    raise RuntimeError("EPUB container does not declare a rootfile")


def _epub_ordered_content_paths(archive: ZipFile, opf_path: str) -> list[str]:
    try:
        opf_xml = archive.read(opf_path)
    except KeyError as exc:
        raise RuntimeError("EPUB package document is missing") from exc
    root = ElementTree.fromstring(opf_xml)
    manifest: dict[str, str] = {}
    spine_ids: list[str] = []
    for element in root.iter():
        local_name = _xml_local_name(element.tag)
        if local_name == "item":
            item_id = (element.attrib.get("id") or "").strip()
            href = (element.attrib.get("href") or "").strip()
            media_type = (element.attrib.get("media-type") or "").strip().lower()
            if item_id and href and media_type in {
                "application/xhtml+xml",
                "text/html",
            }:
                manifest[item_id] = href
        elif local_name == "itemref":
            idref = (element.attrib.get("idref") or "").strip()
            if idref:
                spine_ids.append(idref)

    base_dir = dirname(opf_path)
    ordered = [
        normpath(f"{base_dir}/{manifest[idref]}" if base_dir else manifest[idref])
        for idref in spine_ids
        if idref in manifest
    ]
    if ordered:
        return ordered
    return [
        name
        for name in archive.namelist()
        if name.lower().endswith((".xhtml", ".html", ".htm"))
    ]


def _extract_epub_text(path: Path) -> str:
    return _extract_epub_document(path).text


def _first_text_line(text: str) -> str:
    for line in text.splitlines():
        candidate = " ".join(line.split())
        if candidate:
            return candidate
    return ""


def _extract_epub_document(path: Path) -> ExtractedDocument:
    base_metadata = _base_document_metadata(path, "epub")
    try:
        with ZipFile(path) as archive:
            opf_path = _epub_rootfile_path(archive)
            content_paths = _epub_ordered_content_paths(archive, opf_path)
            segments: list[ExtractedSegment] = []
            for index, content_path in enumerate(content_paths):
                try:
                    text = _strip_html_text(archive.read(content_path))
                except KeyError:
                    continue
                title = _first_text_line(text)
                metadata = {
                    **base_metadata,
                    "spine_index": index,
                    "href": content_path,
                }
                if title:
                    metadata["chapter_title"] = title
                    metadata["section_path"] = [title]
                segment = _segment(text, metadata)
                if segment:
                    segments.append(segment)
    except BadZipFile as exc:
        raise RuntimeError("EPUB file is not a valid zip archive") from exc
    except ElementTree.ParseError as exc:
        raise RuntimeError("EPUB metadata could not be parsed") from exc
    text = "\n\n".join(segment.content for segment in segments).strip()
    metadata = dict(base_metadata)
    metadata["section_count"] = len(segments)
    return ExtractedDocument(text=text, segments=segments, metadata=metadata)


def _extract_pdf_document(path: Path) -> ExtractedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF text extraction requires pypdf to be installed") from exc
    reader = PdfReader(str(path))
    page_count = len(reader.pages)
    base_metadata = {
        **_base_document_metadata(path, "pdf"),
        "page_count": page_count,
    }
    segments: list[ExtractedSegment] = []
    for index, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        segment = _segment(
            text,
            {
                **base_metadata,
                "page_number": index + 1,
            },
        )
        if segment:
            segments.append(segment)
    return ExtractedDocument(
        text="\n\n".join(segment.content for segment in segments).strip(),
        segments=segments,
        metadata=base_metadata,
    )


def _extract_docx_document(path: Path) -> ExtractedDocument:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX text extraction requires python-docx to be installed") from exc

    document = Document(str(path))
    base_metadata = _base_document_metadata(path, "docx")
    segments: list[ExtractedSegment] = []
    current_lines: list[str] = []
    section_path: list[str] = []
    current_section_path: list[str] = []

    def flush() -> None:
        nonlocal current_lines, current_section_path
        segment = _segment(
            "\n".join(current_lines),
            _heading_metadata(base_metadata, current_section_path),
        )
        if segment:
            segments.append(segment)
        current_lines = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (getattr(paragraph.style, "name", "") or "").strip().lower()
        match = re.match(r"heading\s+([1-6])$", style_name)
        if match:
            flush()
            level = int(match.group(1))
            section_path = section_path[: level - 1] + [text]
            current_section_path = list(section_path)
            current_lines = [text]
        else:
            current_lines.append(text)

    flush()
    full_text = "\n".join(
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ).strip()
    if not segments:
        return _single_segment_document(full_text, base_metadata)
    metadata = dict(base_metadata)
    metadata["section_count"] = len(segments)
    return ExtractedDocument(text=full_text, segments=segments, metadata=metadata)


def _extract_document_from_path(path: Path) -> ExtractedDocument:
    suffix = path.suffix.lower()
    data = path.read_bytes()
    if suffix in TEXT_EXTENSIONS:
        text = _decode_text(data)
        base_metadata = _base_document_metadata(
            path,
            "markdown" if suffix in MARKDOWN_EXTENSIONS else suffix.lstrip(".") or "text",
        )
        if suffix in MARKDOWN_EXTENSIONS:
            return _split_markdown_document(text, base_metadata)
        return _single_segment_document(text, base_metadata)
    if suffix == ".pdf":
        return _extract_pdf_document(path)
    if suffix == ".docx":
        return _extract_docx_document(path)
    if suffix == ".epub":
        return _extract_epub_document(path)
    raise RuntimeError(f"Text extraction is not supported for '{suffix or 'unknown'}' files")


def _extract_text_from_path(path: Path) -> str:
    return _extract_document_from_path(path).text


@dataclass
class UploadRecord:
    upload_id: str
    filename: str
    stored_path: str
    content_type: str
    size_bytes: int
    created_at: str
    kind: str
    extracted_text_path: str | None = None
    extracted_manifest_path: str | None = None
    extracted_text_status: str = "not_requested"
    extracted_text_error: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)

    @classmethod
    def from_dict(cls, payload: dict[str, Any] | None) -> "UploadRecord":
        data = payload if isinstance(payload, dict) else {}
        return cls(
            upload_id=str(data.get("upload_id") or "").strip(),
            filename=str(data.get("filename") or "").strip() or "upload.bin",
            stored_path=str(data.get("stored_path") or "").strip(),
            content_type=str(data.get("content_type") or "").strip()
            or "application/octet-stream",
            size_bytes=max(0, int(data.get("size_bytes") or 0)),
            created_at=str(data.get("created_at") or "").strip() or _utcnow_iso(),
            kind=str(data.get("kind") or "").strip() or "binary",
            extracted_text_path=str(data.get("extracted_text_path") or "").strip() or None,
            extracted_manifest_path=str(data.get("extracted_manifest_path") or "").strip()
            or None,
            extracted_text_status=str(data.get("extracted_text_status") or "").strip()
            or "not_requested",
            extracted_text_error=str(data.get("extracted_text_error") or "").strip() or None,
            metadata=data.get("metadata") if isinstance(data.get("metadata"), dict) else {},
        )


class UploadStore:
    def __init__(self, root_dir: str | None = None):
        self.root_dir = Path((root_dir or "").strip() or "kg_agent_uploads")
        self._lock = asyncio.Lock()

    async def save_upload(
        self,
        *,
        filename: str,
        content: bytes,
        content_type: str | None = None,
    ) -> UploadRecord:
        upload_id = uuid.uuid4().hex
        safe_name = _safe_filename(filename)
        resolved_content_type = _guess_content_type(safe_name, content_type)
        kind = _guess_kind(safe_name, resolved_content_type)
        created_at = _utcnow_iso()

        upload_dir = self.root_dir / upload_id
        file_path = upload_dir / safe_name
        metadata_path = upload_dir / "metadata.json"

        record = UploadRecord(
            upload_id=upload_id,
            filename=safe_name,
            stored_path=str(file_path),
            content_type=resolved_content_type,
            size_bytes=len(content),
            created_at=created_at,
            kind=kind,
            metadata={
                "original_filename": filename,
            },
        )

        async with self._lock:
            await asyncio.to_thread(upload_dir.mkdir, parents=True, exist_ok=True)
            await asyncio.to_thread(file_path.write_bytes, content)
            await asyncio.to_thread(
                metadata_path.write_text,
                json.dumps(record.to_dict(), ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
        return record

    async def get_upload(self, upload_id: str) -> UploadRecord | None:
        upload_dir = self.root_dir / (upload_id or "").strip()
        metadata_path = upload_dir / "metadata.json"
        if not metadata_path.exists():
            return None
        try:
            payload = await asyncio.to_thread(metadata_path.read_text, encoding="utf-8")
        except OSError:
            return None
        try:
            parsed = json.loads(payload)
        except json.JSONDecodeError:
            return None
        if not isinstance(parsed, dict):
            return None
        return UploadRecord.from_dict(parsed)

    async def read_bytes(self, upload_id: str) -> bytes:
        record = await self.get_upload(upload_id)
        if record is None:
            raise FileNotFoundError(f"Upload '{upload_id}' not found")
        return await asyncio.to_thread(Path(record.stored_path).read_bytes)

    async def ensure_text_extract(self, upload_id: str) -> UploadRecord:
        async with self._lock:
            record = await self.get_upload(upload_id)
            if record is None:
                raise FileNotFoundError(f"Upload '{upload_id}' not found")

            if (
                record.extracted_text_status == "ready"
                and record.extracted_text_path
                and Path(record.extracted_text_path).exists()
                and record.extracted_manifest_path
                and Path(record.extracted_manifest_path).exists()
            ):
                return record

            try:
                extracted_document = await asyncio.to_thread(
                    _extract_document_from_path,
                    Path(record.stored_path),
                )
                text_path = self.root_dir / record.upload_id / "extracted.txt"
                manifest_path = self.root_dir / record.upload_id / "extracted.json"
                await asyncio.to_thread(
                    text_path.write_text,
                    extracted_document.text,
                    encoding="utf-8",
                )
                await asyncio.to_thread(
                    manifest_path.write_text,
                    json.dumps(
                        extracted_document.to_dict(),
                        ensure_ascii=False,
                        indent=2,
                    ),
                    encoding="utf-8",
                )
                record.extracted_text_path = str(text_path)
                record.extracted_manifest_path = str(manifest_path)
                record.extracted_text_status = "ready"
                record.extracted_text_error = None
            except Exception as exc:
                record.extracted_text_path = None
                record.extracted_manifest_path = None
                record.extracted_text_status = "unsupported"
                record.extracted_text_error = str(exc)

            await self._write_metadata(record)
            return record

    async def read_extracted_text(self, upload_id: str) -> str:
        record = await self.ensure_text_extract(upload_id)
        if record.extracted_text_status != "ready" or not record.extracted_text_path:
            raise RuntimeError(
                record.extracted_text_error
                or f"Upload '{upload_id}' does not have extracted text"
            )
        return await asyncio.to_thread(
            Path(record.extracted_text_path).read_text,
            encoding="utf-8",
        )

    async def read_extracted_document(self, upload_id: str) -> ExtractedDocument:
        record = await self.ensure_text_extract(upload_id)
        if record.extracted_text_status != "ready":
            raise RuntimeError(
                record.extracted_text_error
                or f"Upload '{upload_id}' does not have extracted text"
            )
        if record.extracted_manifest_path and Path(record.extracted_manifest_path).exists():
            payload = await asyncio.to_thread(
                Path(record.extracted_manifest_path).read_text,
                encoding="utf-8",
            )
            try:
                parsed = json.loads(payload)
            except json.JSONDecodeError as exc:
                raise RuntimeError(
                    f"Upload '{upload_id}' extracted manifest is invalid"
                ) from exc
            return ExtractedDocument.from_dict(parsed)
        text = await self.read_extracted_text(upload_id)
        metadata = _base_document_metadata(Path(record.stored_path))
        metadata["source_filename"] = record.filename
        return _single_segment_document(text, metadata)

    async def build_attachment_context(
        self,
        upload_ids: list[str],
        *,
        max_chars_per_attachment: int = 6000,
    ) -> dict[str, Any]:
        attachments: list[dict[str, Any]] = []
        text_blocks: list[str] = []
        unsupported_multimodal = False

        for raw_upload_id in upload_ids:
            upload_id = (raw_upload_id or "").strip()
            if not upload_id:
                continue
            record = await self.get_upload(upload_id)
            if record is None:
                attachments.append(
                    {
                        "upload_id": upload_id,
                        "status": "missing",
                    }
                )
                continue

            item = {
                "upload_id": record.upload_id,
                "filename": record.filename,
                "stored_path": record.stored_path,
                "content_type": record.content_type,
                "kind": record.kind,
                "size_bytes": record.size_bytes,
                "created_at": record.created_at,
                "status": "uploaded",
            }
            if record.extracted_text_path:
                item["extracted_text_path"] = record.extracted_text_path
            if record.extracted_manifest_path:
                item["extracted_manifest_path"] = record.extracted_manifest_path
            if record.metadata:
                item["metadata"] = dict(record.metadata)

            if record.kind in {"text", "document"}:
                resolved = await self.ensure_text_extract(record.upload_id)
                item["text_extract_status"] = resolved.extracted_text_status
                if resolved.extracted_text_status == "ready":
                    extracted_text = await self.read_extracted_text(record.upload_id)
                    preview = extracted_text[:max_chars_per_attachment].strip()
                    if preview:
                        text_blocks.append(
                            f"[附件: {record.filename}]\n{preview}"
                        )
                        item["status"] = "ready"
                    else:
                        item["status"] = "empty_text"
                else:
                    item["status"] = resolved.extracted_text_status
                    item["error"] = resolved.extracted_text_error
            elif record.kind == "image":
                item["status"] = "unsupported_multimodal"
                unsupported_multimodal = True
            else:
                item["status"] = "unsupported_type"
                item["error"] = "当前对话附件仅支持可抽取文本的文档和文本文件。"
            attachments.append(item)

        prompt = ""
        if text_blocks:
            prompt = "\n\n[附件上下文]\n" + "\n\n".join(text_blocks)
        return {
            "attachments": attachments,
            "prompt": prompt,
            "unsupported_multimodal": unsupported_multimodal,
        }

    async def _write_metadata(self, record: UploadRecord) -> None:
        upload_dir = self.root_dir / record.upload_id
        upload_dir.mkdir(parents=True, exist_ok=True)
        metadata_path = upload_dir / "metadata.json"
        await asyncio.to_thread(
            metadata_path.write_text,
            json.dumps(record.to_dict(), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
